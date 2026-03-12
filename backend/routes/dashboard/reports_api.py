"""
Reports API — List and generate reports from the report framework.

Generates real CSV, PDF, and DOCX files from MongoDB data, saves them
to the crm_documents collection via the Documents module.

Security:
- Full tenant isolation via verify_business_access
- All queries scoped to business_id
- Reports are READ-ONLY — never modify data
"""

from fastapi import APIRouter, HTTPException, Depends, Query
from fastapi.responses import JSONResponse
from database import get_database
from middleware.tenant import verify_business_access, TenantContext
from agent.reports import list_reports, get_report, REPORT_REGISTRY
from datetime import datetime, date, timedelta
from typing import Optional
from bson import ObjectId
import csv
import io
import logging

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/reports", tags=["reports"])


def _bid_values(business_id: str):
    """Return list of possible business_id formats for querying."""
    vals = [business_id]
    if ObjectId.is_valid(business_id):
        vals.append(ObjectId(business_id))
    return vals


# ═══════════════════════════════════════
# LIST AVAILABLE REPORTS
# ═══════════════════════════════════════

@router.get("/business/{business_id}")
async def list_available_reports(
    business_id: str,
    category: Optional[str] = Query(None),
    tenant: TenantContext = Depends(verify_business_access),
):
    """List all available report types for this business."""
    reports = list_reports(category)
    return {"reports": reports, "total": len(reports)}


# ═══════════════════════════════════════
# GENERATE A REPORT
# ═══════════════════════════════════════

@router.post("/business/{business_id}/generate")
async def generate_report(
    business_id: str,
    body: dict,
    tenant: TenantContext = Depends(verify_business_access),
):
    """
    Generate a report file and save it to Documents.

    Body:
        report_id: str — which report to generate
        format: str — csv | pdf | docx
        date_from: str (optional) — ISO date
        date_to: str (optional) — ISO date
        params: dict (optional) — extra query params (status, staff, etc.)
    """
    db = get_database()
    bid = tenant.business_id
    bids = _bid_values(bid)

    report_id = body.get("report_id")
    output_format = body.get("format", "csv")
    date_from = body.get("date_from")
    date_to = body.get("date_to")
    extra_params = body.get("params", {})

    if not report_id:
        raise HTTPException(400, "report_id is required")

    report_def = get_report(report_id)
    if not report_def:
        raise HTTPException(404, f"Report '{report_id}' not found")

    if output_format not in report_def["formats"]:
        raise HTTPException(400, f"Format '{output_format}' not supported for this report. Available: {report_def['formats']}")

    # Build query params
    params = {**extra_params}
    if report_def["requires_date_range"]:
        if not date_from:
            date_from = (date.today() - timedelta(days=report_def["default_period_days"])).isoformat()
        if not date_to:
            date_to = date.today().isoformat()
        params["date_from"] = date_from
        params["date_to"] = date_to

    # Run the query builder
    try:
        query = await report_def["query_builder"](db, bid, params)
    except Exception as e:
        logger.error(f"Query builder failed for {report_id}: {e}")
        raise HTTPException(500, f"Failed to build query: {str(e)}")

    # Handle businessId/business_id variants in the query
    if "business_id" in query:
        query.pop("business_id")
        biz_filter = {"$or": [{"businessId": {"$in": bids}}, {"business_id": {"$in": bids}}]}
        # If query already has $or (e.g. staff filter), use $and to combine
        if "$or" in query:
            existing_or = query.pop("$or")
            query["$and"] = [biz_filter, {"$or": existing_or}]
        else:
            query.update(biz_filter)

    # Fetch data
    collection = db[report_def["collection"]]
    sort_args = []
    if report_def.get("sort_field"):
        sort_args = [(report_def["sort_field"], report_def["sort_direction"])]

    cursor = collection.find(query)
    if sort_args:
        cursor = cursor.sort(sort_args)
    cursor = cursor.limit(report_def["max_rows"])

    data = await cursor.to_list(length=report_def["max_rows"])

    if not data:
        return JSONResponse(
            status_code=200,
            content={"success": False, "message": "No data found for the selected period", "count": 0}
        )

    # Build summary if available
    summary = None
    if report_def.get("summary_builder"):
        try:
            summary = await report_def["summary_builder"](db, bid, data)
        except Exception as e:
            logger.warning(f"Summary builder failed for {report_id}: {e}")

    # Get business name for the report header
    business = await db.businesses.find_one({"_id": bid})
    if not business and ObjectId.is_valid(bid):
        business = await db.businesses.find_one({"_id": ObjectId(bid)})
    business_name = business.get("name", "Business") if business else "Business"

    # Generate the file
    fields = report_def["fields"]
    report_name = report_def["name"]

    # Build date label
    date_label = ""
    if date_from and date_to:
        date_label = f"{date_from} to {date_to}"
    elif date_from:
        date_label = f"From {date_from}"

    full_name = f"{report_name}"
    if date_label:
        full_name += f" — {date_label}"

    if output_format == "csv":
        file_bytes = _generate_csv(data, fields, summary)
    elif output_format == "pdf":
        file_bytes = _generate_pdf(data, fields, summary, report_name, business_name, date_label)
    elif output_format == "docx":
        file_bytes = _generate_docx(data, fields, summary, report_name, business_name, date_label)
    else:
        raise HTTPException(400, f"Unsupported format: {output_format}")

    # Save to Documents
    from routes.dashboard.documents import save_document

    doc = await save_document(
        db=db,
        business_id=bid,
        name=full_name,
        file_bytes=file_bytes,
        format=output_format,
        type="report" if report_def["category"] != "exports" else "export",
        tag=_map_category_to_tag(report_def["category"]),
        created_by=tenant.user_email or "Unknown",
        created_by_type="user",
        report_id=report_id,
    )

    return {
        "success": True,
        "document": doc,
        "rows": len(data),
        "summary": summary,
    }


# ═══════════════════════════════════════
# FILE GENERATORS
# ═══════════════════════════════════════

def _get_field_value(record: dict, key: str):
    """Extract a value from a record, supporting nested keys like 'customer.email'."""
    if "." in key:
        parts = key.split(".")
        val = record
        for part in parts:
            if isinstance(val, dict):
                val = val.get(part, "")
            else:
                return ""
        return val

    val = record.get(key, "")

    # Handle service name which can be a dict
    if key == "service_name" and isinstance(val, dict):
        val = val.get("name", str(val))

    # Handle ObjectId
    if isinstance(val, ObjectId):
        val = str(val)

    return val


def _format_field_value(val, fmt: str) -> str:
    """Format a value according to field format type."""
    if val is None or val == "":
        return ""

    if fmt == "currency":
        try:
            return f"£{float(val):,.2f}"
        except (ValueError, TypeError):
            return str(val)
    elif fmt == "date":
        if isinstance(val, str) and len(val) >= 10:
            return val[:10]
        return str(val)
    elif fmt == "datetime":
        if isinstance(val, str) and len(val) >= 16:
            return val[:16].replace("T", " ")
        return str(val)
    elif fmt == "number":
        try:
            return str(int(float(val)))
        except (ValueError, TypeError):
            return str(val)
    elif fmt == "percent":
        try:
            return f"{float(val):.0f}%"
        except (ValueError, TypeError):
            return str(val)
    elif fmt == "status":
        return str(val).replace("_", " ").title()
    elif fmt == "list":
        if isinstance(val, list):
            return ", ".join(str(v) for v in val)
        return str(val)
    else:
        return str(val)


def _generate_csv(data: list, fields: list, summary: dict = None) -> bytes:
    """Generate a CSV file from report data."""
    output = io.StringIO()
    writer = csv.writer(output)

    # Header row
    writer.writerow([f["label"] for f in fields])

    # Data rows
    for record in data:
        row = []
        for f in fields:
            val = _get_field_value(record, f["key"])
            row.append(_format_field_value(val, f["format"]))
        writer.writerow(row)

    # Summary rows
    if summary:
        writer.writerow([])
        writer.writerow(["--- Summary ---"])
        for k, v in summary.items():
            writer.writerow([k, str(v)])

    return output.getvalue().encode("utf-8-sig")  # BOM for Excel compatibility


def _generate_pdf(data: list, fields: list, summary: dict, title: str, business_name: str, date_label: str) -> bytes:
    """Generate a branded PDF report. Uses reportlab."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import mm
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, HRFlowable
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_LEFT, TA_RIGHT, TA_CENTER
    except ImportError:
        raise HTTPException(500, "PDF generation requires reportlab. Install with: pip install reportlab")

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=A4,
        topMargin=20 * mm, bottomMargin=20 * mm,
        leftMargin=15 * mm, rightMargin=15 * mm,
    )

    styles = getSampleStyleSheet()

    # Custom styles — ReeveOS branded
    brand_black = colors.HexColor("#111111")
    brand_gold = colors.HexColor("#C9A84C")
    grey_text = colors.HexColor("#6B7280")
    border_color = colors.HexColor("#E5E7EB")
    header_bg = colors.HexColor("#F9FAFB")

    title_style = ParagraphStyle("ReportTitle", parent=styles["Heading1"], fontSize=18, textColor=brand_black, spaceAfter=4, fontName="Helvetica-Bold")
    subtitle_style = ParagraphStyle("ReportSubtitle", parent=styles["Normal"], fontSize=9, textColor=grey_text, spaceAfter=12)
    header_style = ParagraphStyle("SectionHeader", parent=styles["Heading2"], fontSize=11, textColor=brand_black, spaceBefore=16, spaceAfter=8, fontName="Helvetica-Bold")
    body_style = ParagraphStyle("BodyText", parent=styles["Normal"], fontSize=8, textColor=brand_black)
    summary_label = ParagraphStyle("SummaryLabel", parent=styles["Normal"], fontSize=9, textColor=grey_text)
    summary_value = ParagraphStyle("SummaryValue", parent=styles["Normal"], fontSize=11, textColor=brand_black, fontName="Helvetica-Bold")
    footer_style = ParagraphStyle("Footer", parent=styles["Normal"], fontSize=7, textColor=grey_text, alignment=TA_CENTER)

    elements = []

    # Header: R. mark + business name
    header_data = [
        [Paragraph("<b><font color='#C9A84C' size='16'>R.</font></b>", styles["Normal"]),
         Paragraph(f"<font size='8' color='#6B7280'>{business_name}</font>", ParagraphStyle("RightAlign", parent=styles["Normal"], alignment=TA_RIGHT))]
    ]
    header_table = Table(header_data, colWidths=[doc.width * 0.5, doc.width * 0.5])
    header_table.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
    ]))
    elements.append(header_table)

    # Gold line
    elements.append(HRFlowable(width="100%", thickness=2, color=brand_gold, spaceAfter=12))

    # Title
    elements.append(Paragraph(title, title_style))
    meta_parts = []
    if date_label:
        meta_parts.append(date_label)
    meta_parts.append(f"Generated {datetime.utcnow().strftime('%d %b %Y %H:%M')}")
    meta_parts.append(f"{len(data)} records")
    elements.append(Paragraph(" &middot; ".join(meta_parts), subtitle_style))

    # Summary cards (if present)
    if summary:
        elements.append(Paragraph("Summary", header_style))
        summary_data = []
        summary_row = []
        for i, (k, v) in enumerate(summary.items()):
            cell_content = Paragraph(f"<font size='7' color='#6B7280'>{k}</font><br/><font size='11'><b>{v}</b></font>", styles["Normal"])
            summary_row.append(cell_content)
            if len(summary_row) == 4 or i == len(summary) - 1:
                while len(summary_row) < 4:
                    summary_row.append("")
                summary_data.append(summary_row)
                summary_row = []

        if summary_data:
            col_w = doc.width / 4
            summary_table = Table(summary_data, colWidths=[col_w] * 4)
            summary_table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, -1), header_bg),
                ("BOX", (0, 0), (-1, -1), 0.5, border_color),
                ("INNERGRID", (0, 0), (-1, -1), 0.5, border_color),
                ("TOPPADDING", (0, 0), (-1, -1), 8),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]))
            elements.append(summary_table)
            elements.append(Spacer(1, 12))

    # Data table
    elements.append(Paragraph("Data", header_style))

    # Limit visible columns for PDF readability (max 7)
    visible_fields = fields[:7]

    # Header row
    table_header = [Paragraph(f"<b><font size='7' color='#6B7280'>{f['label']}</font></b>", styles["Normal"]) for f in visible_fields]
    table_data = [table_header]

    # Data rows (limit to 200 for PDF)
    for record in data[:200]:
        row = []
        for f in visible_fields:
            val = _get_field_value(record, f["key"])
            formatted = _format_field_value(val, f["format"])
            row.append(Paragraph(f"<font size='7'>{formatted}</font>", styles["Normal"]))
        table_data.append(row)

    if len(data) > 200:
        table_data.append([Paragraph(f"<font size='7' color='#6B7280'>... and {len(data) - 200} more rows (see CSV export for full data)</font>", styles["Normal"])] + [""] * (len(visible_fields) - 1))

    # Calculate column widths
    n_cols = len(visible_fields)
    col_widths = [doc.width / n_cols] * n_cols

    data_table = Table(table_data, colWidths=col_widths, repeatRows=1)
    data_table.setStyle(TableStyle([
        # Header
        ("BACKGROUND", (0, 0), (-1, 0), header_bg),
        ("TEXTCOLOR", (0, 0), (-1, 0), grey_text),
        ("FONTSIZE", (0, 0), (-1, 0), 7),
        # All cells
        ("BOX", (0, 0), (-1, -1), 0.5, border_color),
        ("INNERGRID", (0, 0), (-1, -1), 0.25, border_color),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        # Alternating rows
        *[("BACKGROUND", (0, i), (-1, i), colors.HexColor("#FAFAFA")) for i in range(2, len(table_data), 2)],
    ]))
    elements.append(data_table)

    # Footer
    elements.append(Spacer(1, 20))
    elements.append(HRFlowable(width="100%", thickness=0.5, color=border_color, spaceAfter=8))
    elements.append(Paragraph(f"ReeveOS &middot; {business_name} &middot; Confidential", footer_style))

    doc.build(elements)
    return buffer.getvalue()


def _generate_docx(data: list, fields: list, summary: dict, title: str, business_name: str, date_label: str) -> bytes:
    """Generate a branded DOCX report. Uses python-docx."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor, Cm, Emu
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
    except ImportError:
        raise HTTPException(500, "DOCX generation requires python-docx. Install with: pip install python-docx")

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # Brand colours
    gold = RGBColor(0xC9, 0xA8, 0x4C)
    black = RGBColor(0x11, 0x11, 0x11)
    grey = RGBColor(0x6B, 0x72, 0x80)
    light_grey = RGBColor(0x9C, 0xA3, 0xAF)

    # Header: R. mark
    header_para = doc.add_paragraph()
    run = header_para.add_run("R.")
    run.font.size = Pt(24)
    run.font.color.rgb = gold
    run.font.bold = True

    # Business name + date
    meta_para = doc.add_paragraph()
    run = meta_para.add_run(business_name)
    run.font.size = Pt(9)
    run.font.color.rgb = grey
    if date_label:
        run = meta_para.add_run(f"  •  {date_label}")
        run.font.size = Pt(9)
        run.font.color.rgb = light_grey

    # Title
    title_para = doc.add_paragraph()
    run = title_para.add_run(title)
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = black

    # Generated date
    gen_para = doc.add_paragraph()
    run = gen_para.add_run(f"Generated {datetime.utcnow().strftime('%d %b %Y %H:%M')}  •  {len(data)} records")
    run.font.size = Pt(8)
    run.font.color.rgb = light_grey

    # Summary
    if summary:
        doc.add_paragraph()
        summary_heading = doc.add_paragraph()
        run = summary_heading.add_run("Summary")
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = black

        summary_table = doc.add_table(rows=1, cols=len(summary))
        summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for i, (k, v) in enumerate(summary.items()):
            cell = summary_table.rows[0].cells[i]
            # Label
            p = cell.paragraphs[0]
            run = p.add_run(k)
            run.font.size = Pt(8)
            run.font.color.rgb = grey
            # Value
            p2 = cell.add_paragraph()
            run = p2.add_run(str(v))
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = black

            # Cell shading
            shading = cell._element.get_or_add_tcPr()
            shading_elm = shading.makeelement(qn("w:shd"), {
                qn("w:val"): "clear",
                qn("w:color"): "auto",
                qn("w:fill"): "F9FAFB",
            })
            shading.append(shading_elm)

    # Data table
    doc.add_paragraph()
    data_heading = doc.add_paragraph()
    run = data_heading.add_run("Data")
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = black

    visible_fields = fields[:8]  # Max 8 columns for readability
    table = doc.add_table(rows=1, cols=len(visible_fields))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    header_row = table.rows[0]
    for i, f in enumerate(visible_fields):
        cell = header_row.cells[i]
        p = cell.paragraphs[0]
        run = p.add_run(f["label"])
        run.font.size = Pt(8)
        run.font.bold = True
        run.font.color.rgb = grey

        shading = cell._element.get_or_add_tcPr()
        shading_elm = shading.makeelement(qn("w:shd"), {
            qn("w:val"): "clear",
            qn("w:color"): "auto",
            qn("w:fill"): "F3F4F6",
        })
        shading.append(shading_elm)

    # Data rows (limit 500 for docx)
    for idx, record in enumerate(data[:500]):
        row = table.add_row()
        for i, f in enumerate(visible_fields):
            val = _get_field_value(record, f["key"])
            formatted = _format_field_value(val, f["format"])
            cell = row.cells[i]
            p = cell.paragraphs[0]
            run = p.add_run(formatted)
            run.font.size = Pt(8)
            run.font.color.rgb = black

            # Alternating row shading
            if idx % 2 == 1:
                shading = cell._element.get_or_add_tcPr()
                shading_elm = shading.makeelement(qn("w:shd"), {
                    qn("w:val"): "clear",
                    qn("w:color"): "auto",
                    qn("w:fill"): "FAFAFA",
                })
                shading.append(shading_elm)

    if len(data) > 500:
        row = table.add_row()
        cell = row.cells[0]
        p = cell.paragraphs[0]
        run = p.add_run(f"... and {len(data) - 500} more rows (see CSV export for full data)")
        run.font.size = Pt(8)
        run.font.color.rgb = light_grey

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(f"ReeveOS  •  {business_name}  •  Confidential")
    run.font.size = Pt(7)
    run.font.color.rgb = light_grey

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _map_category_to_tag(category: str) -> str:
    """Map report framework categories to document tags."""
    mapping = {
        "reports": "bookings",
        "exports": "bookings",
        "financial": "financial",
        "forms": "forms",
    }
    return mapping.get(category, category)
